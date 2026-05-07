import docx
path = r'D:/graduationProject/campus-forum/论文结构参考_业务域重构.docx'
doc = docx.Document(path)

# Validation logic
print("--- Check Chapters ---")
# The user wants "第4章/第6章/第7章"
# Based on earlier results, we see "第四章", "第六章", "第七章" in the text dump
chapters_mapping = {
    "第4章": "第四章",
    "第6章": "第六章",
    "第7章": "第七章"
}
for name, target in chapters_mapping.items():
    found = any(target in p.text for p in doc.paragraphs)
    print(f"{name}: {'Found' if found else 'Not Found'}")

print("--- Check 6.1~6.8 ---")
found_6 = []
for i in range(1, 9):
    title = f"6.{i}"
    found_text = "Not Found"
    for p in doc.paragraphs:
        # Check for the actual section headers (not the TOC)
        if p.text.startswith(title) and ("域" in p.text or "模块" in p.text):
            found_text = p.text.strip()
            # We prefer the "域" version from the latter part of the doc as it's the "重构" part
            if "域" in p.text:
                found_text = p.text.strip()
                break
    print(f"{title}: {found_text}")

print("--- Check Table after 7.3 ---")
def find_table_between(doc, start_text, end_text):
    start_idx = -1
    end_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if start_text in p.text:
            start_idx = i
        if end_text in p.text and start_idx != -1:
            end_idx = i
            break
    if start_idx == -1 or end_idx == -1: return None
    body = doc.element.body
    children = list(body)
    start_el = doc.paragraphs[start_idx]._element
    end_el = doc.paragraphs[end_idx]._element
    start_pos = children.index(start_el)
    end_pos = children.index(end_el)
    for pos in range(start_pos + 1, end_pos):
        if children[pos].tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == children[pos]:
                    return table
    return None

table = find_table_between(doc, "7.3 测试功能", "7.4 测试结果")
if table:
    print(f"Rows: {len(table.rows)}")
    print(f"Cols: {len(table.columns)}")
else:
    print("Table not found")
