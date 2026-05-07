import docx
from docx.shared import Inches
import os

path = r'D:/graduationProject/campus-forum/论文结构参考_业务域重构.docx'
doc = docx.Document(path)

def find_table_between(doc, start_text, end_text):
    start_found = False
    for i, p in enumerate(doc.paragraphs):
        if start_text in p.text:
            start_found = True
            # Look for table after this paragraph but before end_text
            # Tables are in doc.tables, but paragraphs and tables are in doc.element.body
            # A more robust way to find the index:
            pass
    
    # Simpler approach: find the paragraph with start_text, then the next table
    start_idx = -1
    end_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if start_text in p.text:
            start_idx = i
        if end_text in p.text and start_idx != -1:
            end_idx = i
            break
            
    # Find table between these paragraphs
    # In python-docx, doc.tables is a list. We need to find which table is between start_idx and end_idx
    # Body elements can be paragraphs or tables.
    body = doc.element.body
    children = list(body)
    
    start_el = doc.paragraphs[start_idx]._element
    end_el = doc.paragraphs[end_idx]._element
    
    start_pos = children.index(start_el)
    end_pos = children.index(end_el)
    
    for pos in range(start_pos + 1, end_pos):
        if children[pos].tag.endswith('tbl'):
            # Found a table. Now map element to doc.tables
            for table in doc.tables:
                if table._element == children[pos]:
                    return table, pos
    return None, -1

start_txt = "7.3 测试功能"
end_txt = "7.4 测试结果"
table, pos = find_table_between(doc, start_txt, end_txt)

if table:
    if len(table.rows) <= 1:
        # Delete table
        parent = table._element.getparent()
        parent.remove(table._element)
        # We'll insert a new one at the same position or after start_txt
        table = None

if table is None:
    # Find start_txt paragraph again to insert after it
    p_start = None
    for p in doc.paragraphs:
        if start_txt in p.text:
            p_start = p
            break
    
    # Insert table after p_start
    new_tbl = doc.add_table(rows=1, cols=7)
    new_tbl.style = 'Table Grid'
    hdr_cells = new_tbl.rows[0].cells
    headers = ['测试编号', '业务域', '测试项', '前置条件', '步骤', '期望结果', '实际结果']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    
    data = [
        ["TC-UA-01", "用户账号", "登录功能", "账号已注册", "输入账号密码点击登录", "进入首页", "符合预期"],
        ["TC-CC-01", "核心内容", "发布帖子", "已登录", "填写标题内容并提交", "列表显示新帖", "符合预期"],
        ["TC-TO-01", "话题运营", "创建话题", "普通用户", "进入话题页创建", "话题创建成功", "符合预期"],
        ["TC-ACT-01", "活动组织", "报名活动", "活动报名中", "点击报名按钮", "显示报名成功", "符合预期"],
        ["TC-LF-01", "生活服务", "发布二手", "已实名", "上传商品信息", "商品成功上架", "符合预期"],
        ["TC-HELP-01", "互助寻物", "发布寻物", "有遗失物", "填写详情发布", "寻物贴可见", "符合预期"],
        ["TC-IM-01", "即时通讯", "发送消息", "双方互为好友", "打开私聊发消息", "对方收到消息", "符合预期"],
        ["TC-GOV-01", "政务治理", "投诉举报", "内容违规", "点击举报并说明理由", "系统受理举报", "符合预期"],
    ]
    
    for row_data in data:
        row_cells = new_tbl.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val

    # Move new table to correct position (after p_start)
    p_start._element.addnext(new_tbl._element)

doc.save(path)

# Validation
doc = docx.Document(path)
print("--- Check Chapters ---")
chapters = ["第4章", "第6章", "第7章"]
for c in chapters:
    found = any(c in p.text for p in doc.paragraphs)
    print(f"{c}: {'Found' if found else 'Not Found'}")

print("--- Check 6.1~6.8 ---")
for i in range(1, 9):
    title = f"6.{i}"
    found_text = "Not Found"
    for p in doc.paragraphs:
        if p.text.startswith(title):
            found_text = p.text.strip()
            break
    print(f"{title}: {found_text}")

print("--- Check Table after 7.3 ---")
table, _ = find_table_between(doc, start_txt, end_txt)
if table:
    print(f"Rows: {len(table.rows)}")
    print(f"Cols: {len(table.columns)}")
else:
    print("Table not found between 7.3 and 7.4")
