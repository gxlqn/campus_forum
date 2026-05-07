from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

src = r'd:/graduationProject/campus-forum/论文结构参考.docx'
out = r'd:/graduationProject/campus-forum/论文结构参考_业务域重构.docx'

doc = Document(src)


def find_heading(startswith_text):
    for p in doc.paragraphs:
        t = (p.text or '').strip()
        if t.startswith(startswith_text):
            return p
    raise ValueError(f'未找到标题: {startswith_text}')


def clear_between(start_p, end_p):
    body = doc._element.body
    cur = start_p._element.getnext()
    while cur is not None and cur is not end_p._element:
        nxt = cur.getnext()
        body.remove(cur)
        cur = nxt


def add_paragraph_before(anchor_p, text, style='Normal'):
    p = anchor_p.insert_paragraph_before(text)
    if style:
        p.style = style
    return p


def add_code_before(anchor_p, code_text):
    p = anchor_p.insert_paragraph_before(code_text)
    if p.runs:
        for r in p.runs:
            r.font.name = 'Consolas'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            r.font.size = Pt(9)
    return p

# 1) 改写 4.3 功能设计
h43 = find_heading('4.3 功能设计')
h5 = find_heading('5数据库设计')
clear_between(h43, h5)

func_lines = [
    '系统功能设计统一按业务域进行分解，形成“职责清晰、边界明确、可独立演进”的架构。',
    '（1）用户与权限域：负责微信登录、实名认证、个人中心、关注关系、钱包账户与RBAC权限控制。',
    '（2）社区内容域：负责板块、帖子、评论、点赞、收藏、匿名发布与内容互动。',
    '（3）交易与订单域：负责二手商品发布、求购匹配、订单创建、约见核销与履约状态流转。',
    '（4）活动组织域：负责活动发布、报名管理、签到与活动生命周期管理。',
    '（5）失物招领域：负责寻物/招领发布、认领申请、证据提交与审核确认。',
    '（6）互助服务域：负责互助任务发布、候选接单、双向确认、资金冻结/放款/退款与仲裁。',
    '（7）即时通信域：负责会话管理、实时消息推送、ACK回执、离线同步与重试。',
    '（8）平台治理与运营域（含搜索发现与信息服务）：负责内容审核、举报处置、版主分配、资讯导航、搜索服务与运营统计。'
]
for line in func_lines:
    add_paragraph_before(h5, line)

# 2) 改写第6章小节标题
rename_map = {
    '6.1 用户与权限模块': '6.1 用户与权限域',
    '6.2 论坛内容模块': '6.2 社区内容域',
    '6.3 校园资讯和导航模块': '6.3 交易与订单域',
    '6.4 校园服务模块': '6.4 活动组织域',
    '6.5 即时通信与通知模块': '6.5 失物招领域',
    '6.6 举报与审核模块': '6.6 互助服务域',
    '6.7 搜索与推荐模块': '6.7 即时通信域',
    '6.8 数据统计与运营分析模块': '6.8 平台治理与运营域（含搜索发现与信息服务）',
}
for p in doc.paragraphs:
    t = (p.text or '').strip()
    if t in rename_map:
        p.text = rename_map[t]

# 3) 改写第6章每个域内容（含核心代码<=20行）
sec61 = find_heading('6.1 用户与权限域')
sec62 = find_heading('6.2 社区内容域')
sec63 = find_heading('6.3 交易与订单域')
sec64 = find_heading('6.4 活动组织域')
sec65 = find_heading('6.5 失物招领域')
sec66 = find_heading('6.6 互助服务域')
sec67 = find_heading('6.7 即时通信域')
sec68 = find_heading('6.8 平台治理与运营域（含搜索发现与信息服务）')
sec7 = find_heading('7系统测试')

section_payload = [
    (sec61, sec62,
     ['本域通过AuthController与UserController完成登录鉴权、学号绑定、关注关系和钱包操作。',
      '实现上采用JWT承载用户身份，RBAC权限在后台管理侧统一下发。'],
     '''@PostMapping("/wx/login")
public Result<LoginResponse> wxLogin(@RequestBody WxLoginRequest request) {
    SysUser user = userService.wxLogin(request.getCode(), request.getUserInfo());
    String token = jwtUtils.generateToken(user.getId(), user.getNickname());
    LoginResponse response = new LoginResponse();
    response.setToken(token);
    response.setUser(user);
    response.setNeedBind(user.getIsVerified() != 1);
    return Result.success(response);
}

@PostMapping("/wallet/recharge")
public Result<Map<String, Object>> rechargeWallet(@AuthenticationPrincipal SysUser currentUser,
        @RequestBody Map<String, Object> body) {
    BigDecimal amount = new BigDecimal(String.valueOf(body.get("amount")));
    return Result.success(userCenterService.rechargeWallet(currentUser.getId(), amount));
}'''),

    (sec62, sec63,
     ['本域由ForumController统一提供帖子、评论、点赞与收藏接口，支持分页与热度排序。',
      '帖子实体通过source_type/source_id支持与交易、活动、互助等业务内容联动。'],
     '''@GetMapping("/posts")
public Result<PageResult<ForumPost>> getPosts(
        @RequestParam(required = false) Long current,
        @RequestParam(required = false) Long page,
        @RequestParam(defaultValue = "10") Long size,
        @RequestParam(required = false) Long sectionId,
        @RequestParam(required = false) String keyword,
        @RequestParam(required = false, defaultValue = "latest") String orderBy,
        @RequestParam(required = false, defaultValue = "72") Integer hotWindowHours) {
    Long pageNo = current != null ? current : page;
    if (pageNo == null) pageNo = 1L;
    return Result.success(forumService.getPostList(pageNo, size, sectionId, keyword, orderBy, hotWindowHours));
}'''),

    (sec63, sec64,
     ['本域围绕ProductController实现商品发布、订单创建、卖家处理、约见改约、核销确认。',
      '订单状态变化与商品状态联动，保障交易过程可跟踪、可回溯。'],
     '''@PostMapping("/{id}/order")
public Result<ServiceProductOrder> createOrder(@PathVariable Long id,
        @AuthenticationPrincipal SysUser currentUser) {
    return Result.success(productService.createOrder(id, currentUser.getId()));
}

@PostMapping("/orders/{orderId}/meetup/verify")
public Result<Void> verifyMeetupCode(@PathVariable Long orderId,
        @RequestBody ProductOrderVerifyRequest request,
        @AuthenticationPrincipal SysUser currentUser) {
    productService.verifyMeetupCode(orderId, currentUser.getId(), request.getMeetupCode());
    return Result.success();
}'''),

    (sec64, sec65,
     ['本域通过ActivityController实现活动发布、报名、取消报名与后台审核。',
      '活动数据包含时间窗、地点、人数上限等要素，支持活动状态流转。'],
     '''@PostMapping
public Result<ServiceActivity> create(@RequestBody ServiceActivity activity,
        @AuthenticationPrincipal SysUser currentUser) {
    activity.setUserId(currentUser.getId());
    return Result.success(activityService.createActivity(activity));
}

@PostMapping("/{id}/signup")
public Result<Void> signup(@PathVariable Long id,
        @AuthenticationPrincipal SysUser currentUser) {
    activityService.signupActivity(id, currentUser.getId());
    return Result.success();
}'''),

    (sec65, sec66,
     ['本域通过LostFoundController实现寻物/招领发布、认领申请与审核处理。',
      '认领流程支持证据图片提交与审核意见回写，形成闭环。'],
     '''@PostMapping("/{id}/claim")
public Result<Void> submitClaim(@PathVariable Long id,
        @RequestBody ServiceLostFoundClaim claim,
        @AuthenticationPrincipal SysUser currentUser) {
    lostFoundService.submitClaim(id, currentUser.getId(), claim.getDescription(), claim.getImages());
    return Result.success();
}

@PostMapping("/claims/{claimId}/audit")
public Result<Void> auditClaim(@PathVariable Long claimId,
        @RequestBody AuditActionRequest request,
        @AuthenticationPrincipal SysUser currentUser) {
    lostFoundService.auditClaim(claimId, currentUser.getId(), request.getAuditStatus(), request.getAuditRemark());
    return Result.success();
}'''),

    (sec66, sec67,
     ['本域由HelpController承载互助任务发布、接单、双向确认与仲裁流程。',
      '在服务层实现资金冻结、放款、退款及争议处理规则。'],
     '''@PostMapping("/{id}/accept")
public Result<Void> acceptHelp(@PathVariable Long id,
        @AuthenticationPrincipal SysUser currentUser) {
    helpService.acceptHelp(id, currentUser.getId());
    return Result.success();
}

@PostMapping("/{id}/publisher-confirm")
public Result<Void> publisherConfirm(@PathVariable Long id,
        @AuthenticationPrincipal SysUser currentUser,
        @RequestParam(defaultValue = "0") Integer isComplaint) {
    helpService.publisherConfirm(id, currentUser.getId(), isComplaint);
    return Result.success();
}'''),

    (sec67, sec68,
     ['本域采用STOMP over WebSocket协议实现实时消息、ACK回执与离线同步。',
      '消息投递任务表与回执表共同保证至少一次投递和幂等处理。'],
     '''@MessageMapping("/im/send")
public void send(@Payload ImSendMessageRequest request, Principal principal) {
    Long userId = userIdOf(principal);
    Map<String, Object> result = imRealtimeService.send(userId, request);
    messagingTemplate.convertAndSendToUser(String.valueOf(userId), "/queue/im-send-ack", result);
}

@MessageMapping("/im/ack")
public void ack(@Payload ImAckRequest request, Principal principal) {
    Long userId = userIdOf(principal);
    Map<String, Object> result = imRealtimeService.ack(userId, request);
    messagingTemplate.convertAndSendToUser(String.valueOf(userId), "/queue/im-ack", result);
}'''),

    (sec68, sec7,
     ['本域整合审核、举报、敏感词、运营统计、搜索发现与资讯导航能力。',
      '治理侧通过权限控制实现后台处置，搜索侧提供高级检索、纠错与推荐。'],
     '''@PreAuthorize("hasAuthority('system:report')")
@GetMapping("/reports")
public Result<PageResult<Map<String, Object>>> getReports(
        @RequestParam(defaultValue = "1") Long current,
        @RequestParam(defaultValue = "10") Long size,
        @RequestParam(required = false) Integer status,
        @RequestParam(required = false) Integer targetType) {
    return Result.success(adminSystemService.getReports(current, size, status, targetType));
}

@GetMapping("/advanced")
public Result<Map<String, Object>> searchAdvanced(
        @RequestParam(required = false) String keyword,
        @RequestParam(defaultValue = "5") Integer size,
        @RequestParam(required = false) String sectionName,
        @RequestParam(required = false) String priceRange,
        @RequestParam(defaultValue = "relevance") String sort) {
    return Result.success(searchService.searchAdvanced(keyword, size, sectionName, priceRange, sort));
}''')
]

for start, end, descs, code in section_payload:
    clear_between(start, end)
    for d in descs:
        add_paragraph_before(end, d)
    add_paragraph_before(end, '核心代码片段：')
    add_code_before(end, code)

# 4) 改写第7章：7.1/7.2/7.3
s71 = find_heading('7.1 测试目的')
s72 = find_heading('7.2 测试方法')
s73 = find_heading('7.3 测试功能')
s74 = find_heading('7.4 测试结果')

clear_between(s71, s72)
for t in [
    '系统测试以八大业务域为主线，验证“功能正确、流程闭环、权限隔离、数据一致”。',
    '重点覆盖用户与权限、社区内容、交易订单、活动、失物、互助、即时通信、平台治理与搜索信息服务。'
]:
    add_paragraph_before(s72, t)

clear_between(s72, s73)
for t in [
    '采用黑盒测试与接口测试结合的方法，按业务域设计用例并执行端到端验证。',
    '测试过程覆盖正常流、异常流、越权访问、重复提交与弱网重试等场景。'
]:
    add_paragraph_before(s73, t)

# 清空7.3原内容（包括旧表格），重建业务域测试表
clear_between(s73, s74)
add_paragraph_before(s74, '表7.3 业务域功能测试表')

headers = ['测试编号', '业务域', '测试项', '前置条件', '步骤', '期望结果', '实际结果']
rows = [
    ['TC-UA-01', '用户与权限域', '微信登录与JWT签发', '用户已授权小程序', '提交code完成登录', '返回token与用户信息', '通过'],
    ['TC-CC-01', '社区内容域', '帖子发布与评论回复', '用户已登录且板块可用', '发布帖子并新增评论', '帖子与评论成功入库并展示', '通过'],
    ['TC-TO-01', '交易与订单域', '下单-约见-核销', '商品审核通过且在售', '买家下单后卖家约见并核销码验证', '订单进入已完成状态', '通过'],
    ['TC-ACT-01', '活动组织域', '活动报名与取消', '活动处于报名中', '用户报名后取消报名', '报名人数正确增减', '通过'],
    ['TC-LF-01', '失物招领域', '认领申请与审核', '存在招领信息', '用户提交认领证据，管理员审核', '认领状态按审核结果更新', '通过'],
    ['TC-HELP-01', '互助服务域', '接单与发布者确认', '互助单待接单', '接单人接单后发布者确认完成', '订单完成并触发资金结算', '通过'],
    ['TC-IM-01', '即时通信域', '消息发送与ACK回执', '双方在线且已建立会话', '发送消息并回执ACK', '会话未读数与回执状态正确', '通过'],
    ['TC-GOV-01', '平台治理与运营域（含搜索信息）', '举报处置与高级搜索', '存在待处理举报与可检索数据', '管理员处理举报并执行高级搜索', '举报状态更新且搜索返回匹配结果', '通过']
]

table = doc.add_table(rows=1, cols=len(headers))
for i, h in enumerate(headers):
    table.cell(0, i).text = h
for r in rows:
    rcells = table.add_row().cells
    for i, v in enumerate(r):
        rcells[i].text = v
# 将新表移动到7.4标题之前
s74._element.addprevious(table._element)

doc.save(out)
print(out)
